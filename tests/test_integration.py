import pytest
from pathlib import Path
from docx import Document
from markdocx import MarkDocx
from . import get_test_file, create_temp_file, cleanup_temp_file

def test_full_document_conversion():
    """测试完整文档转换"""
    # 使用测试资源中的完整测试文件
    test_md = get_test_file("test.md")
    
    converter = MarkDocx()
    output_path = create_temp_file("", ".docx")
    
    try:
        # 转换文档
        result = converter.convert(str(test_md), str(output_path))
        
        # 验证输出文件存在
        assert result.exists()
        assert result.stat().st_size > 0
        
        # 验证文档内容
        doc = Document(output_path)
        
        # 验证标题
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) >= 2, "Missing headings"
        assert headings[0].text.strip() == "Test Document"
        
        # 验证文本格式
        found_formatting = False
        for p in doc.paragraphs:
            text = p.text
            if "bold" in text:
                found_formatting = any(run.bold for run in p.runs)
                if found_formatting:
                    break
        assert found_formatting, "Text formatting not properly applied"
        
        # 验证列表
        list_items = [p for p in doc.paragraphs if p.style.name.startswith('List')]
        assert len(list_items) >= 4, "Missing list items"
        
        # 验证代码块
        found_code = False
        for p in doc.paragraphs:
            if "hello()" in p.text:
                found_code = True
                break
        assert found_code, "Code block not found"
        
        # 验证表格
        assert len(doc.tables) > 0, "Table not found"
        table = doc.tables[0]
        assert table.rows[0].cells[0].text.strip() == "Header 1"
        
        # 验证图片
        found_image = False
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                found_image = True
                break
        assert found_image, "Image not found"
        
        # 验证链接
        found_link = False
        for rel in doc.part.rels.values():
            if "hyperlink" in rel.reltype.lower():
                found_link = True
                break
        assert found_link, "Hyperlink not found"
        
    finally:
        cleanup_temp_file(output_path)

def test_custom_style_conversion():
    """测试使用自定义样式的文档转换"""
    test_md = get_test_file("test.md")
    
    # 自定义样式
    style_config = {
        "normal": {
            "font": {
                "default": "Arial",
                "east-asia": "微软雅黑",
                "size": 11,
                "color": "000000"
            },
            "line-spacing": 1.15,
            "space": {
                "before": 5,
                "after": 5
            }
        },
        "heading_1": {
            "font": {
                "size": 20,
                "extra": ["bold"]
            }
        }
    }
    
    converter = MarkDocx(style_config=style_config)
    output_path = create_temp_file("", ".docx")
    
    try:
        result = converter.convert(str(test_md), str(output_path))
        
        # 验证样式应用
        doc = Document(output_path)
        
        # 验证标题样式
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        heading_style = doc.styles['Heading1']
        font_size = heading_style._element.get_or_add_rPr().sz_val
        assert font_size == 20 * 2 * 12700  # Word中的字号单位转换
        assert heading_style.font.bold
        
        # 验证正文样式
        normal_style = doc.styles['Normal']
        if normal_style:
            font = normal_style.font
            assert font.name == 'Arial'
            
    finally:
        cleanup_temp_file(output_path)

def test_batch_conversion():
    """测试批量文档转换"""
    # 创建多个测试文件
    test_files = []
    for i in range(3):
        content = f"""# Document {i+1}
        
This is test document {i+1}.

- Item 1
- Item 2

```python
print("Hello from doc {i+1}")
```
"""
        test_files.append(create_temp_file(content, ".md"))
    
    converter = MarkDocx()
    output_files = []
    
    try:
        # 批量转换
        for test_file in test_files:
            output_path = create_temp_file("", ".docx")
            output_files.append(output_path)
            result = converter.convert(str(test_file), str(output_path))
            
            # 验证每个输出文件
            assert result.exists()
            assert result.stat().st_size > 0
            
            # 验证基本内容
            doc = Document(output_path)
            assert len(doc.paragraphs) > 0
            assert any(p.style.name.startswith('Heading') for p in doc.paragraphs)
            
    finally:
        # 清理所有临时文件
        for file in test_files + output_files:
            cleanup_temp_file(file)

if __name__ == "__main__":
    pytest.main([__file__]) 