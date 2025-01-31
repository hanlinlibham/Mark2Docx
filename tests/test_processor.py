import pytest
from pathlib import Path
from bs4 import BeautifulSoup
from docx import Document
from markdocx.core.processor import DocxProcessor
from . import get_test_file, create_temp_file, cleanup_temp_file

def test_process_link():
    """测试超链接处理"""
    processor = DocxProcessor()
    html = """
    <body>
        <p>This is a <a href="https://example.com">link</a> in text.</p>
    </body>
    """
    soup = BeautifulSoup(html, 'html.parser')
    
    # 使用临时文件进行测试
    output_path = create_temp_file("", ".docx")
    try:
        processor.process(soup, str(output_path))
        
        # 验证生成的文档中的超链接
        doc = Document(output_path)
        found_link = False
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                if hasattr(run, '_element') and run._element.hyperlink:
                    found_link = True
                    assert run._element.hyperlink.get('href') == "https://example.com"
        assert found_link, "No hyperlink found in document"
    finally:
        cleanup_temp_file(output_path)

def test_process_image():
    """测试图片处理"""
    processor = DocxProcessor()
    test_image = get_test_file("test.png")
    html = f"""
    <body>
        <p><img src="{test_image}" alt="Test Image"></p>
    </body>
    """
    soup = BeautifulSoup(html, 'html.parser')
    
    output_path = create_temp_file("", ".docx")
    try:
        processor.process(soup, str(output_path))
        
        # 验证生成的文档中的图片
        doc = Document(output_path)
        found_image = False
        for rel in doc.part.rels.values():
            if "image" in rel.reltype:
                found_image = True
                break
        assert found_image, "No image found in document"
    finally:
        cleanup_temp_file(output_path)

def test_process_table():
    """测试表格处理"""
    processor = DocxProcessor()
    html = """
    <body>
        <table>
            <tr><th>Header 1</th><th>Header 2</th></tr>
            <tr><td>Cell 1</td><td>Cell 2</td></tr>
        </table>
    </body>
    """
    soup = BeautifulSoup(html, 'html.parser')
    
    output_path = create_temp_file("", ".docx")
    try:
        processor.process(soup, str(output_path))
        
        # 验证生成的文档中的表格
        doc = Document(output_path)
        assert len(doc.tables) > 0, "No table found in document"
        table = doc.tables[0]
        assert table.rows[0].cells[0].text.strip() == "Header 1"
        assert table.rows[1].cells[1].text.strip() == "Cell 2"
    finally:
        cleanup_temp_file(output_path)

def test_process_list():
    """测试列表处理"""
    processor = DocxProcessor()
    html = """
    <body>
        <ul>
            <li>Item 1</li>
            <li>Item 2
                <ul>
                    <li>Sub item 1</li>
                    <li>Sub item 2</li>
                </ul>
            </li>
        </ul>
    </body>
    """
    soup = BeautifulSoup(html, 'html.parser')
    
    output_path = create_temp_file("", ".docx")
    try:
        processor.process(soup, str(output_path))
        
        # 验证生成的文档中的列表
        doc = Document(output_path)
        list_paragraphs = [p for p in doc.paragraphs if p.style.name.startswith('List')]
        assert len(list_paragraphs) >= 4, "List items not properly processed"
    finally:
        cleanup_temp_file(output_path)

def test_process_heading():
    """测试标题处理"""
    processor = DocxProcessor()
    html = """
    <body>
        <h1>Heading 1</h1>
        <h2>Heading 2</h2>
        <h3>Heading 3</h3>
    </body>
    """
    soup = BeautifulSoup(html, 'html.parser')
    
    output_path = create_temp_file("", ".docx")
    try:
        processor.process(soup, str(output_path))
        
        # 验证生成的文档中的标题
        doc = Document(output_path)
        headings = [p for p in doc.paragraphs if p.style.name.startswith('Heading')]
        assert len(headings) == 3, "Headings not properly processed"
        assert headings[0].text.strip() == "Heading 1"
        assert headings[1].text.strip() == "Heading 2"
        assert headings[2].text.strip() == "Heading 3"
    finally:
        cleanup_temp_file(output_path)

def test_progress_callback():
    """测试进度回调"""
    progress_values = []
    messages = []
    
    def callback(progress: float, message: str):
        progress_values.append(progress)
        messages.append(message)
    
    processor = DocxProcessor()
    processor.set_progress_callback(callback)
    
    html = """
    <body>
        <h1>Title</h1>
        <p>Paragraph 1</p>
        <p>Paragraph 2</p>
    </body>
    """
    soup = BeautifulSoup(html, 'html.parser')
    
    output_path = create_temp_file("", ".docx")
    try:
        processor.process(soup, str(output_path))
        
        # 验证进度回调
        assert len(progress_values) > 0
        assert progress_values[-1] == 1.0  # 最后一个进度值应该是1.0
        assert "completed" in messages[-1].lower()  # 最后一个消息应该包含"completed"
    finally:
        cleanup_temp_file(output_path) 