# noinspection PyProtectedMember
#
import io
import os
import re
import docx
from urllib.request import urlopen

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import *
from docx.oxml import parse_xml, OxmlElement
from docx.oxml.ns import qn, nsdecls
from docx.oxml.shared import OxmlElement, qn
from docx.shape import InlineShape
from docx.shared import Inches, RGBColor, Pt
from docx.table import Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from docx.opc.constants import RELATIONSHIP_TYPE
from requests import HTTPError

from ..provider.docx_plus import add_hyperlink
from ..provider.style_manager import StyleManager
from ..utils.style_enum import MDX_STYLE

debug_state: bool = False
auto_open: bool = True
show_image_desc: bool = True  # 是否显示图片的描述，即 `![desc](src/img)` 中 desc的内容


def debug(*args):
    print(*args) if debug_state else None


class DocxProcessor:
    def __init__(self, style_conf: dict):
        self.document = Document()
        if style_conf is not None:
            StyleManager(self.document, style_conf).init_styles()

    # h1, h2, ...
    def add_heading(self, content: str, tag: str):
        level: int = int(tag.__getitem__(1))
        p = self.document.add_paragraph(content, style="Heading%d" % level)
        return p

    # noinspection PyMethodMayBeStatic
    def add_run(self, p: Paragraph, content: str, char_style: str = "plain"):
        # fixme 行内的样式超过一个的句子会被忽略，如：
        # <u>**又加粗又*斜体*又下划线**</u>
        debug("[%s]:" % char_style, content)
        run = p.add_run(content)

        # 不应当使用形如 run.bold = (char_style=="strong") 的方式
        # 因为没有显式加粗，不意味着整体段落不加粗。
        if char_style == "strong":
            run.bold = True
        if char_style == "em":
            run.italic = True
        if char_style == "u":
            run.underline = True
        if char_style == "strike":
            run.font.strike = True
        if char_style == "sub":
            run.font.subscript = True
        if char_style == "sup":
            run.font.superscript = True
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW if char_style == "highlight" else None

        # if char_style == "code":
        #     run.font.name = "Consolas"

    def add_code_block(self, pre_tag):
        """处理代码块"""
        # 获取代码内容
        code_content = ""
        if hasattr(pre_tag, 'code') and pre_tag.code is not None:
            code_content = pre_tag.code.string
        else:
            code_content = pre_tag.string
            
        if code_content:
            # 去除多余的换行符
            code_content = code_content.strip()
            # 创建代码块段落
            p = self.document.add_paragraph(style=MDX_STYLE.PLAIN_TEXT)
            run = p.add_run(code_content)
            run.font.name = "Consolas"
            # 设置代码块的格式
            p.paragraph_format.first_line_indent = 0
            p.paragraph_format.left_indent = Pt(20)  # 左缩进
            p.paragraph_format.space_before = Pt(10)  # 段前间距
            p.paragraph_format.space_after = Pt(10)  # 段后间距

    def add_picture(self, img_tag):
        p: Paragraph = self.document.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        run: Run = p.add_run()
        p.paragraph_format.first_line_indent = 0

        img_src: str
        scale: float = 100  # 优先级最高，单位 %
        width_px: int = 100
        height_px: int = 100

        # 设置宽度
        if img_tag.get("style"):
            style_content: str = img_tag["style"]
            img_attr: list = style_content.strip().split(";")
            # print(img_attr)
            attr: str
            for attr in img_attr:
                if attr.find("width") != -1:
                    # TODO 处理 style 中的宽度和高度属性
                    width_px = int(re.findall(r"\d+", attr)[0])
                if attr.find("height") != -1:
                    height_px = int(re.findall(r"\d+", attr)[0])
                if attr.find("zoom") != -1:
                    scale = int(re.findall(r"\d+", attr)[0])

        if img_tag["src"] != "":
            img_src = img_tag["src"]
            # 网络图片
            if img_src.startswith("http://") or img_src.startswith("https://"):
                print("[IMAGE] fetching:", img_src)
                try:
                    image_bytes = urlopen(img_src, timeout=10).read()
                    data_stream = io.BytesIO(image_bytes)
                    run.add_picture(data_stream, width=Inches(5.7 * scale / 100))
                except Exception as e:
                    print("[RESOURCE ERROR]:", e)
            else:
                # 本地图片
                run.add_picture(img_src, width=Inches(5.7 * scale / 100))
        else:
            # 网络图片
            img_src = img_tag["title"]
            print("[IMAGE] fetching:", img_src)
            try:
                image_bytes = urlopen(img_src, timeout=10).read()
                data_stream = io.BytesIO(image_bytes)
                run.add_picture(data_stream, width=Inches(5.7 * scale / 100))
            except Exception as e:
                print("[RESOURCE ERROR]:", e)

        # 如果选择展示图片描述，那么描述会在图片下方显示
        if show_image_desc and img_tag.get("alt"):
            # TODO 图片描述的显示样式
            desc: Paragraph = self.document.add_paragraph(img_tag["alt"], style=MDX_STYLE.CAPTION)
            desc.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            desc.style.font.color.rgb = RGBColor(11, 11, 11)
            desc.style.font.bold = False
            desc.paragraph_format.first_line_indent = 0

    def add_table(self, table_root):
        # 统计列数
        col_count: int = 0
        for col in table_root.thead.tr.contents:
            if col.string != "\n":
                col_count += 1

        table = self.document.add_table(0, col_count, style=MDX_STYLE.TABLE)  # TODO 表格样式

        # 表格头行
        head_row_cells = table.add_row().cells
        i = 0
        for col in table_root.thead.tr.contents:
            if col.string == "\n":
                continue
            head_row_cells[i].paragraphs[0].add_run(col.string).bold = True  # TODO 表内单元格字符样式
            i += 1

        # 数据行
        for tr in table_root.tbody:
            if tr.string == "\n":
                continue
            row_cells = table.add_row().cells
            i = 0
            for td in tr.contents:
                if td.string == "\n":
                    continue
                row_cells[i].text = td.string
                i += 1

    def add_number_list(self, number_list):
        # print(number_list.contents, "\n")
        num: int = 1  # 序号
        for item in number_list.children:
            if item.string == "\n":
                continue
            self.add_paragraph(item, p_style=MDX_STYLE.LIST_NUMBER) \
                .style.paragraph_format.space_after = Pt(1)  # TODO 数字列表样式

            if hasattr(item, "ol") and item.ol is not None:  # 有子序列
                sub_num: int = 1  # 子序号
                for item2 in item.ol.children:
                    if item2.string == "\n":
                        continue
                    self.add_paragraph(item2, prefix="(%d). " % sub_num, p_style=MDX_STYLE.LIST_CONTINUE) \
                        .style.paragraph_format.first_line_indent = 0  # TODO 数字列表样式
                    sub_num += 1
            num += 1

    def add_bullet_list(self, bullet_list):
        # 有可能是TODO list
        text = str(bullet_list.contents[1].string).strip()
        if text.startswith("[ ]") or text.startswith("[x]"):
            self.add_todo_list(bullet_list)
            return
        for item in bullet_list.children:
            text: str = str(item.string)
            if text == "\n":
                continue
            self.add_paragraph(item, p_style=MDX_STYLE.LIST_BULLET) \
                .style.paragraph_format.space_after = Pt(1)  # TODO 无序列表样式 ·• ‣°º৹ ■ ◻ ■ □ ◉◎ ●◌

            if hasattr(item, "ul") and item.ul is not None:  # 有子序列
                for item2 in item.ul.children:
                    if item2.string == "\n":
                        continue
                    # list_para.add_run("   ◉ " + str(item2.string) + "\n")
                    self.add_paragraph(item2, prefix="•  ", p_style=MDX_STYLE.LIST_CONTINUE) \
                        .style.paragraph_format.space_after = Pt(1)  # TODO 数字列表样式

    # 伪TODO list
    def add_todo_list(self, todo_list):
        # list_para.style.font.name = "Consolas"
        for item in todo_list.children:
            if item.string == "\n":
                continue
            text: str = item.string
            list_para = self.document.add_paragraph(style=MDX_STYLE.PLAIN_LIST)
            if text.startswith("[x]"):
                list_para.add_run("[ √ ]").font.name = "Consolas"
                list_para.add_run(text.replace("[x]", " ", 1))
            if text.startswith("[ ]"):
                list_para.add_run("[   ]").font.name = "Consolas"
                list_para.add_run(text.replace("[ ]", " ", 1))

    # 分割线，转换为 Word 中的分页符
    def add_split_line(self):
        self.document.add_page_break()

    # 超链接
    def add_link(self, p: Paragraph, text: str, href: str):
        """添加超链接"""
        debug("[link]:", text, "[href]:", href)
        # 创建超链接
        part = p.part
        r_id = part.relate_to(href, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
        
        # 创建超链接元素
        hyperlink = OxmlElement('w:hyperlink')
        hyperlink.set(qn('r:id'), r_id)
        
        # 创建文本运行元素
        new_run = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        
        # 将超链接添加到段落
        p._p.append(hyperlink)
        
        # 设置超链接样式
        run = p.add_run()
        run.font.color.rgb = RGBColor(0, 0, 255)
        run.font.underline = True

    def add_paragraph(self, children, p_style: str = None, prefix: str = ""):
        if not children:
            return

        if isinstance(children, str):
            p = self.document.add_paragraph(prefix + children)
            if p_style:
                p.style = p_style
            return

        # 处理特殊标签
        if children.name == "pre":
            self.add_code_block(children)
            return
        elif children.name == "img":
            self.add_picture(children)
            return
        elif children.name == "table":
            self.add_table(children)
            return
        elif children.name == "ol":
            self.add_number_list(children)
            return
        elif children.name == "ul":
            if "task-list" in children.get("class", []):
                self.add_todo_list(children)
            else:
                self.add_bullet_list(children)
            return
        elif children.name == "hr":
            self.add_split_line()
            return
        elif children.name == "blockquote":
            self.add_blockquote(children)
            return

        p = self.document.add_paragraph(prefix, style=p_style)
        if type(children) == str:
            p.add_run(children)
            return p
        for elem in children.contents:  # 遍历一个段落内的所有元素
            if elem.name == "a":
                self.add_link(p, elem.string, elem["href"])
            elif elem.name == "img":
                self.add_picture(elem)
            elif elem.name is not None:  # 有字符样式的子串
                self.add_run(p, elem.string, elem.name)
            elif not elem.string == "\n":  # 无字符样式的子串
                self.add_run(p, elem)
        return p

    # from docx.enum.style import WD_STYLE
    def add_blockquote(self, children):
        # TODO 将引用块放在1x1的表格中，优化引用块的显示效果
        #  设置左侧缩进，上下行距
        table: Table = self.document.add_table(0, 1)
        row_cells = table.add_row().cells
        p = row_cells[0].paragraphs[0]

        for child in children.contents:
            if child.string != "\n":
                # self.add_paragraph(p, p_style=MDX_STYLE.BLOCKQUOTE)
                if type(child) == str:
                    p.add_run(child)
                    return p
                for elem in child.contents:  # 遍历一个段落内的所有元素
                    if elem.name == "a":
                        self.add_link(p, elem.string, elem["href"])
                    elif elem.name == "img":
                        self.add_picture(elem)
                    elif elem.name is not None:  # 有字符样式的子串
                        self.add_run(p, elem.string, elem.name)
                    elif not elem.string == "\n":  # 无字符样式的子串
                        self.add_run(p, elem)

        shading_elm_1 = parse_xml(r'<w:shd {} w:fill="efefef"/>'.format(nsdecls('w')))
        table.rows[0].cells[0]._tc.get_or_add_tcPr().append(shading_elm_1)
        # table_format = table.style.paragraph_format

        # 直接操作 Oxml 的方式设置左侧缩进和表格宽度
        # noinspection PyProtectedMember
        tbl_pr = table._element.xpath('w:tblPr')
        # if tbl_pr:
            # 左侧缩进
            # e = OxmlElement('w:tblInd')
            # e.set(qn('w:w'), "300")
            # e.set(qn('w:type'), 'dxa')
            # tbl_pr[0].append(e)
            # 设置表格宽度
            # w = OxmlElement('w:tblW')
            # w.set(qn('w:w'), "4700")
            # w.set(qn('w:type'), "pct")
            # tbl_pr[0].append(w)


    def html2docx(self, html_path: str, docx_path: str):
        # 打开HTML
        with open(html_path, 'r', encoding="UTF-8") as html_file:
            html_str = html_file.read()
        soup = BeautifulSoup(html_str, 'html.parser')
        body_tag = soup.contents[2]
        # 将工作目录切换到指定目录
        os.chdir(os.path.dirname(os.path.abspath(html_path)))
        # 逐个解析标签，并写到word中
        for root in body_tag.children:
            if root.string != "\n":
                # debug("<%s>" % root.name)
                if root.name == "p":  # 普通段落
                    self.add_paragraph(root, p_style=MDX_STYLE.PLAIN_TEXT)
                if root.name == "blockquote":  # 引用块
                    self.add_blockquote(root)
                if root.name == "ol":  # 数字列表
                    self.add_number_list(root)
                if root.name == "ul":  # 无序列表 或 TODO_List
                    self.add_bullet_list(root)
                if root.name == "table":  # 表格
                    self.add_table(root)
                if root.name == "hr":
                    self.add_split_line()
                if root.name == "pre":
                    self.add_code_block(root)
                if root.name == "h1" or root.name == "h2" or \
                        root.name == "h3" or root.name == "h4" or root.name == "h5":
                    self.add_heading(root.string, root.name)

        self.document.save(docx_path)
